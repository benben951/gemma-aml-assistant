"""AML Compliance Assistant - Streamlit Frontend"""

import streamlit as st
from typing import Optional

st.set_page_config(
    page_title="AML Compliance Assistant",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.title("AML Compliance Assistant")
st.markdown("> Safety & Trust Hackathon - Gemma 4 powered compliance assistant")

# Sidebar
with st.sidebar:
    st.header("Config")
    
    backend = st.selectbox(
        "Backend",
        ["ollama", "kaggle", "huggingface"],
        index=0,
    )
    
    enable_thinking = st.checkbox("Thinking Mode", value=True)
    top_k = st.slider("Top K Results", 1, 10, 5)
    
    storage_type = st.selectbox(
        "Vector Store",
        ["memory", "qdrant_local", "qdrant_cloud"],
        index=0,
    )
    
    st.markdown("---")
    
    if "query_count" not in st.session_state:
        st.session_state.query_count = 0
    if "doc_count" not in st.session_state:
        st.session_state.doc_count = 0
    
    st.metric("Queries", st.session_state.query_count)
    st.metric("Documents", st.session_state.doc_count)


@st.cache_resource
def init_gemma_client(backend: str):
    try:
        from src.logic.gemma_client import GemmaClient
        return GemmaClient(backend=backend)
    except Exception as e:
        st.error(f"Failed to init Gemma client: {e}")
        return None


@st.cache_resource
def init_pipeline(storage_type: str):
    try:
        from src.logic.rag_pipeline import RAGPipeline
        return RAGPipeline(storage_type=storage_type)
    except Exception as e:
        st.error(f"Failed to init RAG pipeline: {e}")
        return None


@st.cache_resource
def init_explainability(gemma_client):
    try:
        from src.logic.explainability import ExplainabilityEngine
        return ExplainabilityEngine(gemma_client)
    except Exception as e:
        st.error(f"Failed to init explainability engine: {e}")
        return None


tab1, tab2, tab3 = st.tabs(["Q&A", "Documents", "System"])

# Tab 1: Q&A
with tab1:
    st.header("Compliance Q&A")
    
    question = st.text_area(
        "Ask your question",
        placeholder="e.g., What is AML due diligence?",
        height=100,
    )
    
    col1, col2 = st.columns([1, 4])
    with col1:
        submit_btn = st.button("Query", type="primary")
    with col2:
        st.markdown("*Enable Thinking Mode to see reasoning process*")
    
    if submit_btn and question:
        client = init_gemma_client(backend)
        pipeline = init_pipeline(storage_type)
        
        if client and pipeline:
            with st.spinner("Searching..."):
                try:
                    # Step 1: Retrieve relevant documents
                    search_results = pipeline.retrieve(question, top_k=top_k)
                    
                    if not search_results:
                        st.warning("No relevant documents found. Please upload documents first.")
                    else:
                        # Step 2: Use explainability engine for answer
                        engine = init_explainability(client)
                        result = engine.analyze_with_sources(
                            query=question,
                            search_results=search_results,
                            enable_thinking=enable_thinking,
                        )
                        st.session_state.query_count += 1
                        
                        st.markdown("---")
                        
                        if enable_thinking and result.thinking:
                            with st.expander("Thinking Process", expanded=True):
                                st.markdown(result.thinking)
                        
                        st.subheader("Answer")
                        st.markdown(result.answer)
                        
                        if result.sources:
                            st.subheader("Sources")
                            for i, source in enumerate(result.sources[:3]):
                                st.markdown(f"**[{i+1}] {source.get('source', 'Unknown')}**")
                                excerpt = source.get('excerpt', '')
                                if excerpt:
                                    st.markdown(f"> {excerpt[:200]}...")
                                if source.get('score'):
                                    st.caption(f"Relevance: {source['score']:.2f}")
                        
                        st.subheader("Confidence")
                        confidence = result.confidence
                        st.progress(confidence)
                        
                        if confidence >= 0.7:
                            st.success(f"Confidence: {confidence:.1%} - High reliability")
                        elif confidence >= 0.5:
                            st.warning(f"Confidence: {confidence:.1%} - Verify recommended")
                        else:
                            st.error(f"Confidence: {confidence:.1%} - Low reliability")
                    
                except Exception as e:
                    st.error(f"Query failed: {e}")

# Tab 2: Documents
with tab2:
    st.header("Document Management")
    
    st.markdown("Upload AML regulations. Supports PDF, Word, text files.")
    
    MAX_FILE_SIZE = 10 * 1024 * 1024
    MAX_FILES = 10
    
    uploaded_files = st.file_uploader(
        "Upload",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
    )
    
    if uploaded_files:
        if len(uploaded_files) > MAX_FILES:
            st.error(f"Max {MAX_FILES} files")
            uploaded_files = uploaded_files[:MAX_FILES]
        
        valid_files = []
        for file in uploaded_files:
            if file.size > MAX_FILE_SIZE:
                st.warning(f"{file.name} exceeds 10MB, skipped")
            else:
                valid_files.append(file)
        
        uploaded_files = valid_files
        
        if uploaded_files:
            st.write(f"Uploaded {len(uploaded_files)} files")
            for file in uploaded_files:
                st.write(f"- {file.name} ({file.size / 1024:.1f} KB)")
        
        process_btn = st.button("Process", type="primary")
        
        if process_btn:
            pipeline = init_pipeline(storage_type)
            
            if pipeline:
                with st.spinner("Processing..."):
                    try:
                        from src.data.models import Document
                        
                        documents = []
                        
                        for file in uploaded_files:
                            content = file.read()
                            
                            if file.name.endswith(".txt"):
                                text = content.decode("utf-8")
                            elif file.name.endswith(".pdf"):
                                try:
                                    import fitz
                                    doc = fitz.open(stream=content, filetype="pdf")
                                    text = ""
                                    for page in doc:
                                        text += page.get_text()
                                except ImportError:
                                    text = f"[PDF: {file.name}]"
                            elif file.name.endswith(".docx"):
                                try:
                                    from docx import Document as DocxDocument
                                    doc = DocxDocument(file)
                                    text = "\n".join([para.text for para in doc.paragraphs])
                                except ImportError:
                                    text = f"[Word: {file.name}]"
                            else:
                                text = f"[Unknown: {file.name}]"
                            
                            documents.append(
                                Document(
                                    content=text,
                                    source=file.name,
                                )
                            )
                        
                        count = pipeline.add_documents(documents)
                        st.session_state.doc_count += len(uploaded_files)
                        
                        st.success(f"Processed {len(uploaded_files)} files, {count} chunks")
                        
                    except Exception as e:
                        st.error(f"Failed: {e}")
    
    st.markdown("---")
    st.subheader("Current Collection")
    
    pipeline = init_pipeline(storage_type)
    if pipeline:
        try:
            info = pipeline.get_collection_info()
            st.json(info)
        except Exception:
            st.info("Collection empty")
    
    if st.button("Clear Collection"):
        pipeline = init_pipeline(storage_type)
        if pipeline:
            pipeline.delete_collection()
            st.session_state.doc_count = 0
            st.success("Cleared")
            st.rerun()

# Tab 3: System
with tab3:
    st.header("System Info")
    
    st.subheader("Gemma Client")
    client = init_gemma_client(backend)
    if client:
        st.json({
            "backend": client.backend.value,
            "model": client.model,
        })
    else:
        st.warning("Not initialized")
    
    st.subheader("RAG Pipeline")
    pipeline = init_pipeline(storage_type)
    if pipeline:
        try:
            info = pipeline.get_collection_info()
            st.json(info)
        except Exception:
            st.json({"storage_type": pipeline.storage_type})
    else:
        st.warning("Not initialized")
    
    st.markdown("---")
    st.subheader("Quick Start")
    
    st.markdown("""
1. Select backend in sidebar (Ollama recommended for local)
2. Upload documents in Documents tab
3. Ask questions in Q&A tab

**Safety & Trust Features:**
- Thinking Mode: shows reasoning
- Source citations: traceable answers  
- Confidence score: quality indicator
    """)


st.markdown("---")
st.markdown("""
<div style="text-align: center; color: gray;">
    <small>AML Compliance Assistant | Safety & Trust Hackathon | Gemma 4</small>
</div>
""", unsafe_allow_html=True)